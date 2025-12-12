"""Génère le manuel d'utilisation (DOCX) pour SAE501.

Objectif
- Document orienté « utilisateur / enseignant » : comment lancer et utiliser l'application.
- 3 modes exigés : local, hébergé, VM (git clone).

Sortie
- MANUEL_UTILISATION_SAE501.docx
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "MANUEL_UTILISATION_SAE501.docx"

# Métadonnées (page de garde)
BINOME_ETUDIANT_1 = "Diego Massat"
BINOME_ETUDIANT_2 = "Shun Von Lunen"
GROUPE_NUMERO = "17"
STATUT = "Non-alternants"


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


def add_title_page(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading("MANUEL D’UTILISATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(30)

    subtitle = doc.add_paragraph("SAE501 - Plateforme de recherche et recommandation de séries TV")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Binôme : ").bold = True
    meta.add_run(f"{BINOME_ETUDIANT_1} - {BINOME_ETUDIANT_2}\n")
    meta.add_run("Groupe : ").bold = True
    meta.add_run(f"{GROUPE_NUMERO} ({STATUT})\n")
    meta.add_run("Date : ").bold = True
    meta.add_run(date.today().strftime("%d/%m/%Y") + "\n")
    meta.add_run("URL locale : ").bold = True
    meta.add_run("http://127.0.0.1:5000\n")
    meta.add_run("Technologies : ").bold = True
    meta.add_run("Python, Flask, SQLite, JavaScript (Vanilla)")

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("SOMMAIRE", 1)

    items = [
        "1. Objectif du manuel",
        "2. Prérequis",
        "3. Lancement en local (Windows / PowerShell)",
        "4. Lancement hébergé (Procfile / gunicorn)",
        "5. Lancement via VM (git clone)",
        "6. Utilisation de l’interface web",
        "7. Tests rapides (API)",
        "8. Dépannage (problèmes courants)",
    ]

    for it in items:
        doc.add_paragraph(it, style="List Number")

    doc.add_page_break()


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_section(doc: Document, title: str, paragraphs: list[str]) -> None:
    doc.add_heading(title, 1)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_subsection(doc: Document, title: str, paragraphs: list[str]) -> None:
    doc.add_heading(title, 2)
    for text in paragraphs:
        doc.add_paragraph(text)


def main() -> None:
    doc = Document()
    set_margins(doc)

    add_title_page(doc)
    add_toc(doc)

    # 1. Objectif
    add_section(
        doc,
        "1. OBJECTIF DU MANUEL",
        [
            "Ce document explique comment lancer et utiliser l’application SAE501 (recherche et recommandation de séries TV).",
            "Il est conçu pour un enseignant ou un évaluateur qui souhaite vérifier rapidement le fonctionnement en conditions réelles.",
            "Trois modes de lancement sont décrits : (1) local, (2) hébergé, (3) via machine virtuelle (VM) avec un git clone.",
        ],
    )

    # 2. Prérequis
    add_section(
        doc,
        "2. PRÉREQUIS",
        [
            "Pour exécuter l’application, il faut :",
            "- Python 3.8+ (recommandé : Python 3.11+)",
            "- Git (pour cloner le dépôt, surtout en VM)",
            "- Accès réseau (optionnel) : uniquement si vous souhaitez récupérer des posters externes ; sinon l’application fonctionne avec les données embarquées.",
            "Sur Windows / PowerShell, l’activation du venv peut nécessiter une autorisation d’exécution des scripts (voir section Dépannage).",
        ],
    )

    add_subsection(
        doc,
        "2.1 Dépendances Python",
        [
            "Les dépendances sont listées dans requirements.txt. L’installation se fait via pip.",
            "Le serveur peut être lancé en mode développement (Flask) ou en mode production (gunicorn) selon le contexte.",
        ],
    )

    # 3. Local
    add_section(
        doc,
        "3. LANCEMENT EN LOCAL (WINDOWS / POWERSHELL)",
        [
            "Ce mode est celui à privilégier pour une démonstration rapide sur un poste Windows.",
            "Les commandes ci-dessous supposent que vous êtes dans le dossier du projet SAE501.",
        ],
    )

    add_subsection(
        doc,
        "3.1 Installation (une seule fois)",
        [
            "1) Créer un environnement virtuel Python :",
        ],
    )
    add_code_block(
        doc,
        "python -m venv venv\n",
    )

    add_subsection(doc, "3.2 Activer l’environnement virtuel", ["Dans Windows PowerShell :"])
    add_code_block(doc, ".\\venv\\Scripts\\Activate.ps1\n")

    add_subsection(doc, "3.3 Installer les dépendances", ["Installer les packages Python :"])
    add_code_block(doc, "pip install -r requirements.txt\n")

    add_subsection(
        doc,
        "3.4 Créer/initialiser la base SQLite",
        [
            "Créer la base data/tvseries.db (2-3 minutes selon la machine) :",
        ],
    )
    add_code_block(doc, "python database\\import_data_sqlite.py\n")

    add_subsection(
        doc,
        "3.5 Démarrer le serveur",
        [
            "Lancer le serveur Flask (SQLite) :",
        ],
    )
    add_code_block(doc, "python serve_api_sqlite.py\n")

    add_subsection(
        doc,
        "3.6 Accéder à l’application",
        [
            "Ouvrir un navigateur sur : http://127.0.0.1:5000",
            "Pages utiles : /search.html, /recommendations.html, /series.html, /profile.html",
        ],
    )

    # 4. Hébergé
    add_section(
        doc,
        "4. LANCEMENT HÉBERGÉ (PROCFILE / GUNICORN)",
        [
            "Ce mode correspond à un déploiement sur une plateforme d’hébergement (PaaS) qui fournit une variable d’environnement PORT.",
            "Le projet inclut un Procfile et un script start.sh pour lancer le serveur en gunicorn.",
            "Important : la base SQLite (data/tvseries.db) doit être disponible sur le disque de l’instance. Si l’hébergement est éphémère, la base sera recréée au démarrage.",
        ],
    )

    add_subsection(
        doc,
        "4.1 Commande de démarrage (gunicorn)",
        [
            "Commande de service (équivalente à celle du Procfile) :",
        ],
    )
    add_code_block(doc, "gunicorn serve_api_sqlite:app --bind 0.0.0.0:$PORT\n")

    add_subsection(
        doc,
        "4.2 Initialisation automatique de la base (start.sh)",
        [
            "Le script start.sh vérifie si data/tvseries.db existe, et lance l’import SQLite si nécessaire.",
            "Sur une plateforme Linux, le point d’entrée peut être configuré pour exécuter start.sh.",
        ],
    )
    add_code_block(
        doc,
        "# start.sh (logique)\n" "python database/import_data_sqlite.py  # si DB absente\n" "gunicorn serve_api_sqlite:app --bind 0.0.0.0:$PORT\n",
    )

    add_subsection(
        doc,
        "4.3 Vérification après déploiement",
        [
            "Une fois l’application accessible via une URL publique, vérifier :",
            "- La page web (/) se charge",
            "- La documentation API (/api) répond en JSON",
            "- La recherche renvoie des résultats sur une requête simple",
        ],
    )

    # 5. VM
    add_section(
        doc,
        "5. LANCEMENT VIA VM (GIT CLONE)",
        [
            "Ce mode est prévu pour une démonstration sur une machine virtuelle fournie avec un OS propre.",
            "Le principe est identique au mode local, avec une étape supplémentaire : cloner le dépôt.",
        ],
    )

    add_subsection(doc, "5.1 Cloner le dépôt", ["Dans un terminal, exécuter :"])
    add_code_block(doc, "git clone https://github.com/Diego0777/SAE501.git\ncd SAE501\n")

    add_subsection(
        doc,
        "5.2 Installer et démarrer",
        [
            "Répéter ensuite les étapes du mode local : venv → pip → import SQLite → serveur.",
            "En pratique :",
        ],
    )
    add_code_block(
        doc,
        "python -m venv venv\n"
        ".\\venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt\n"
        "python database\\import_data_sqlite.py\n"
        "python serve_api_sqlite.py\n",
    )

    add_subsection(
        doc,
        "5.3 Validation rapide",
        [
            "Ouvrir http://127.0.0.1:5000 puis effectuer une recherche (ex : crash avion ile) sur /search.html.",
        ],
    )

    # 6. Utilisation UI
    add_section(
        doc,
        "6. UTILISATION DE L’INTERFACE WEB",
        [
            "Une fois le serveur démarré, l’interface web est disponible sur http://127.0.0.1:5000.",
            "Le site est une interface statique (web/) qui appelle l’API REST via fetch.",
        ],
    )

    add_subsection(
        doc,
        "6.1 Parcours recommandé (démonstration en 3 minutes)",
        [
            "1) Aller sur /search.html et rechercher : crash avion ile (Lost doit apparaître très haut).",
            "2) Aller sur /series.html, ouvrir une série, consulter les mots-clés et notes.",
            "3) Aller sur /profile.html, se connecter avec un compte de test puis noter une série.",
            "4) Aller sur /recommendations.html pour afficher les recommandations (popularité ou personnalisées).",
        ],
    )

    add_subsection(
        doc,
        "6.2 Comptes de test",
        [
            "Des comptes de test existent dans la base SQLite importée (ex : alice / bob / charlie…).",
            "Les identifiants exacts peuvent être fournis dans la documentation du projet (GUIDE_DEMARRAGE.md).",
        ],
    )

    # 7. Tests API
    add_section(
        doc,
        "7. TESTS RAPIDES (API)",
        [
            "Les tests ci-dessous permettent de vérifier le cœur fonctionnel sans passer par l’interface web.",
            "Sous Windows PowerShell, on peut utiliser Invoke-RestMethod.",
        ],
    )

    add_subsection(doc, "7.1 Documentation API", ["Vérifier que /api répond :"])
    add_code_block(doc, 'Invoke-RestMethod -Uri "http://127.0.0.1:5000/api"\n')

    add_subsection(doc, "7.2 Recherche (exigence SAE)", ["Exemple :"])
    add_code_block(
        doc,
        'Invoke-RestMethod -Uri "http://127.0.0.1:5000/api/search?q=crash+avion+ile&limit=3"\n',
    )

    add_subsection(doc, "7.3 Recommandations", ["Popularité (top 5) :"])
    add_code_block(
        doc,
        'Invoke-RestMethod -Uri "http://127.0.0.1:5000/api/recommend/popularity?limit=5"\n',
    )

    add_subsection(
        doc,
        "7.4 Tests automatiques",
        [
            "Le projet inclut un script d’intégration SQLite :",
        ],
    )
    add_code_block(doc, "python test_integration_sqlite.py\n")

    # 8. Dépannage
    add_section(
        doc,
        "8. DÉPANNAGE (PROBLÈMES COURANTS)",
        [
            "Cette section liste les problèmes les plus fréquents lors de l’évaluation (poste Windows, VM, hébergement).",
        ],
    )

    add_subsection(
        doc,
        "8.1 PowerShell : activation venv refusée",
        [
            "Symptôme : l’activation .\\venv\\Scripts\\Activate.ps1 est bloquée.",
            "Solution (temporaire, session courante) :",
        ],
    )
    add_code_block(doc, "Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass\n")

    add_subsection(
        doc,
        "8.2 Port 5000 déjà utilisé",
        [
            "Symptôme : le serveur ne démarre pas car le port est occupé.",
            "Solutions : arrêter le processus qui écoute, ou relancer sur un autre port via la variable PORT.",
        ],
    )

    add_subsection(
        doc,
        "8.3 Base SQLite absente",
        [
            "Symptôme : erreurs liées à data/tvseries.db.",
            "Solution : relancer l’import SQLite : python database\\import_data_sqlite.py",
        ],
    )

    doc.save(str(OUT_DOCX))
    print(f"DOCX créé : {OUT_DOCX}")


if __name__ == "__main__":
    main()
